import docx
import os
import re
import pyexcel as pe

def read_file(f_name):
    """
    reading docx files using python-docx package and returning whole text.
    """
    doc = docx.Document(f_name)

    text_in_doc = [doc.paragraphs[i].text for i in range(len(doc.paragraphs))]
    text_s = " ".join(text_in_doc)  # change list to string.
    text_s = re.sub('[^A-Za-z0-9 ]+', '', text_s)
    return text_s.lower()


def analyze(f_name):
    """
    calculate word count and later the respective frequency in a dictionary.
    Returns sorted list with words which have frequency more than 0.001.

    """
    word_list = {}
    words = read_file(f_name)
    for i in words.split():  # creating word count dictionary
        if (i in word_list):
            word_list[i] = word_list[i] + 1
        else:
            word_list[i] = 1

    n = len(words.split())
    for i in word_list:  # creating frequency dictionary
        word_list[i] = round(word_list[i]/n, 4)

    w_table = {i: word_list[i] for i in word_list if word_list[i] > 0.001}
    sorted_values = sorted(w_table.items(), key = lambda x:x[1], reverse = True)
    sorted_list = [[k,v] for k,v in sorted_values]

    return sorted_list

docfile = "ulysses"
data = analyze(docfile+".docx")
dic = {"Word Frequency Stats": data}
pe.save_book_as(bookdict = dic, dest_file_name = docfile+"_word_stat.xlsx")
